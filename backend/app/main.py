from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
from io import BytesIO
import httpx
from dotenv import load_dotenv
import os

from .models import TextAlignment


load_dotenv()
GALILEU_TOKEN = os.getenv("GALILEU_TOKEN")
GALILEU_URL = os.getenv("GALILEU_URL")
MAPBOX_TOKEN = os.getenv("MAPBOX_TOKEN")
POCKETBASE_URL = os.getenv("POCKETBASE_URL")

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def format_paragraph(paragraph, html: str):
    soup = BeautifulSoup(html, "html.parser")
    for element in soup.descendants:
        if element.name is None:
            run = paragraph.add_run(element)
            parent = element.parent
            if parent.find_parent("b") or parent.name == "b":
                run.bold = True
            if parent.find_parent("i") or parent.name == "i":
                run.italic = True
            if parent.find_parent("u") or parent.name == "u":
                run.underline = True
    return paragraph


def align_paragraph(paragraph, alignment: TextAlignment):
    if alignment == TextAlignment.CENTER.value:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == TextAlignment.JUSTIFY.value:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == TextAlignment.RIGHT.value:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == TextAlignment.LEFT.value:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return paragraph


def process_template_data(document: dict):

    if not document.get("data"):
        raise ValueError("No data to process in the template.")

    processed_result = {
        "templateId": document["templateId"],
        "templateName": document["templateName"],
        "pages": document["data"],
    }

    return processed_result


def save_processed_result_to_docx(processed_result: dict):

    document = Document()
    h1_style = document.styles.add_style("H1", WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = "Arial"
    h1_style.font.size = Pt(24)
    h1_style.font.bold = False
    h1_style.font.italic = False

    h2_style = document.styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = "Arial"
    h2_style.font.size = Pt(18)
    h2_style.font.bold = False
    h2_style.font.italic = False

    h3_style = document.styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = "Arial"
    h3_style.font.size = Pt(14)
    h3_style.font.bold = False
    h3_style.font.italic = False

    p_style = document.styles.add_style("P", WD_STYLE_TYPE.PARAGRAPH)
    p_style.font.name = "Arial"
    p_style.font.size = Pt(12)
    p_style.font.bold = False
    p_style.font.italic = False

    for page in processed_result["pages"]:
        rows = page["pageData"]["blocks"]

        for section in document.sections:
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height

            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        for i, row in enumerate(rows):
            print(row)
            if row["type"] == "header":
                text = row["data"]["text"]
                level = int(row["data"]["level"])
                alignment = row["tunes"]["alignmentTune"]["alignment"]

                header = document.add_heading(level=level)

                if level == 1:
                    header.style = document.styles["H1"]
                elif level == 2:
                    header.style = document.styles["H2"]
                elif level == 3:
                    header.style = document.styles["H3"]

                header = align_paragraph(header, alignment)
                header = format_paragraph(header, text)

            elif row["type"] == "paragraph":
                text = row["data"]["text"]
                alignment = row["tunes"]["alignmentTune"]["alignment"]

                paragraph = document.add_paragraph()

                paragraph.style = document.styles["P"]

                paragraph = align_paragraph(paragraph, alignment)
                paragraph = format_paragraph(paragraph, text)

            elif row["type"] == "list":
                items = row["data"]["items"]
                alignment = row["tunes"]["alignmentTune"]["alignment"]

                for item in items:
                    text = item["content"]
                    style = row["data"]["style"]
                    paragraph = document.add_paragraph()

                    if style == "unordered":
                        paragraph.style = document.styles["List Bullet"]
                    elif style == "ordered":
                        paragraph.style = document.styles["List Number"]

                    paragraph = align_paragraph(paragraph, alignment)
                    paragraph = format_paragraph(paragraph, text)

                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(12)

            elif row["type"] == "table":
                table_content = row["data"]["content"]
                have_header = row["data"]["withHeadings"]
                num_rows = len(table_content)
                num_cols = len(table_content[0])

                table = document.add_table(rows=num_rows, cols=num_cols)
                table.style = document.styles["Table Grid"]

                for i, row in enumerate(table_content):
                    for j, cell_text in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = ""

                        parts = [
                            part.strip()
                            for part in cell_text.split("<br>")
                            if part.strip()
                        ]

                        paragraph = cell.paragraphs[0]

                        for idx, part in enumerate(parts):
                            if idx > 0:
                                paragraph.add_run().add_break()
                            run = paragraph.add_run(part)
                            run.font.name = "Arial"
                            run.font.size = Pt(12)

                            if i == 0 and have_header:
                                run.font.bold = True

            elif row["type"] == "image":
                url = row["data"]["file"]["url"]
                caption = row["data"]["caption"]
                response = httpx.get(url)
                response.raise_for_status()

                image_stream = BytesIO(response.content)

                document.add_picture(image_stream, width=Inches(6))

                if caption != "":
                    caption_paragraph = document.add_paragraph(
                        str(caption).replace("<br>", "")
                    )
                    caption_paragraph.style = document.styles["P"]
                    caption_paragraph = align_paragraph(
                        caption_paragraph, TextAlignment.CENTER.value
                    )
                    for run in caption_paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)

            # Handle line breaks
            if (i < len(rows) - 1) and rows[i + 1]["type"] == "breakLine":
                break_line = document.add_paragraph()
                break_line.style = document.styles["P"]
                run = break_line.add_run()
        document.add_page_break()

    doc_stream = BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream


@app.post("/templates/")
async def receive_template_record(record: dict):

    try:
        processed_data = process_template_data(record)

        doc_stream = save_processed_result_to_docx(processed_data)

        return StreamingResponse(
            doc_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=document_report.docx"
            },
        )

    except ValueError as e:
        print(e)
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(e)
        raise HTTPException(
            status_code=500, detail="An error occurred while processing the data."
        )


@app.get("/galileu/{laudo}")
async def get_galileu_data(laudo: int):
    async with httpx.AsyncClient() as client:
        response = await client.get(
            f"{GALILEU_URL}/{laudo}",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {GALILEU_TOKEN}",
            },
        )

    if response.status_code != 200:
        raise HTTPException(
            status_code=response.status_code,
            detail=f"Error fetching data: {response.text}",
        )

    return response.json()


async def upload_image_to_pocketbase(image_bytes: bytes) -> str:
    collection = "images"

    file_tuple = ("image.jpg", image_bytes, "image/jpeg")
    files = {"file": file_tuple}

    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{POCKETBASE_URL}/api/collections/{collection}/records", files=files
        )
    response.raise_for_status()

    record = response.json()
    record_id = record.get("id")
    file_name = record.get("file")

    return f"{POCKETBASE_URL}/api/files/{collection}/{record_id}/{file_name}"


@app.post("/mapbox/")
async def get_mapbox_img(lon: float, lat: float, zoom: float):
    pb_url = (
        f"https://api.mapbox.com/styles/v1/mapbox/satellite-v9/static/"
        f"pin-s+555555({lon},{lat})/{lon},{lat},{zoom},0/512x512"
        f"?access_token={MAPBOX_TOKEN}"
    )

    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(pb_url)
            response.raise_for_status()
    except httpx.RequestError as exc:
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred while requesting the external API: {exc}",
        )
    except httpx.HTTPStatusError as exc:
        raise HTTPException(
            status_code=exc.response.status_code,
            detail=f"External API error: {exc.response.text}",
        )
    data = await upload_image_to_pocketbase(response.content)
    return {"url": data}
